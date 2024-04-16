from docx import Document
import PyPDF2
import pandas as pd


def extract_file_content(file_path):
    """
    处理文件，提取其中的内容。
    """
    if file_path.endswith('.docx'):
        return extract_docx_content(file_path)
    elif file_path.endswith(('.xlsx', '.xls')):
        return extract_xlsx_content(file_path)
    elif file_path.endswith('.pdf'):
        return extract_pdf_content(file_path)
    else:
        raise ValueError(f"{file_path}并非 docx 或 xlsx 或 xls 或 pdf 文件")


def extract_docx_content(docx_path):
    """
    从.docx文件中提取文本和表格内容。
    """
    doc = Document(docx_path)
    content = []

    # 提取文本
    for para in doc.paragraphs:
        content.append(para.text)

    # 提取表格内容
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                content.append(cell.text)

    return '\n'.join(content)


def extract_xlsx_content(xlsx_path):
    """
    从 .xlsx 文件中提取所有工作表的内容，并合并成单个字符串。
    """
    xlsx = pd.ExcelFile(xlsx_path)
    sheet_contents = []

    for sheet_name in xlsx.sheet_names:
        df = pd.read_excel(xlsx, sheet_name)
        # 将 DataFrame 转换为字符串，每行用换行符分隔
        sheet_content = df.to_string(index=False, header=True)
        sheet_contents.append(sheet_content)

    return '\n'.join(sheet_contents)


def extract_pdf_content(pdf_path):
    """
    使用 PyPDF2 库从 PDF 文件中提取文本内容。
    """
    content = []
    with open(pdf_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            content.append(page.extract_text())
    return '\n'.join(content)